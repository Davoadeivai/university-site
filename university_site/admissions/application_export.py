"""خروجی لیست درخواست‌های پذیرش به اکسل، ورد و چاپ."""
from __future__ import annotations

from collections import defaultdict
from io import BytesIO

from django.http import HttpResponse
from django.utils import timezone
from django.utils.html import escape

from core.jalali import format_jalali_datetime


HEADERS = [
    'ردیف',
    'کد رهگیری',
    'نام',
    'نام خانوادگی',
    'نام پدر',
    'کد ملی',
    'موبایل',
    'ایمیل',
    'جنسیت',
    'مقطع',
    'رشته (اولویت ۱)',
    'رشته (اولویت ۲)',
    'گروه آموزشی',
    'دانشکده',
    'شیفت',
    'وضعیت',
    'معدل',
    'تاریخ ثبت',
]


def application_row(index: int, app) -> list:
    major = app.desired_major
    major2 = app.desired_major2
    group = ''
    department = ''
    if major:
        if getattr(major, 'group_id', None) and major.group:
            group = major.group.name
        if getattr(major, 'department_id', None) and major.department:
            department = major.department.name
    created = format_jalali_datetime(app.created_at) if app.created_at else ''
    return [
        index,
        app.tracking_code or '',
        app.first_name or '',
        app.last_name or '',
        app.father_name or '',
        app.national_id or '',
        app.phone or '',
        app.email or '',
        app.get_gender_display() if app.gender else '',
        app.get_degree_display() if app.degree else '',
        major.name if major else '',
        major2.name if major2 else '',
        group,
        department,
        app.get_shift_display() if app.shift else '',
        app.get_status_display() if app.status else '',
        str(app.gpa) if app.gpa is not None else '',
        created,
    ]


def build_rows(applications) -> list[list]:
    return [application_row(i, app) for i, app in enumerate(applications, start=1)]


def build_summary(applications) -> list[tuple[str, int]]:
    """شمارش بر اساس مقطع > گروه > رشته"""
    buckets: dict[str, int] = defaultdict(int)
    for app in applications:
        degree = app.get_degree_display() if app.degree else '—'
        major = app.desired_major
        group = '—'
        major_name = '—'
        if major:
            major_name = major.name
            if getattr(major, 'group_id', None) and major.group:
                group = major.group.name
        key = f'{degree} | {group} | {major_name}'
        buckets[key] += 1
    return sorted(buckets.items(), key=lambda x: (-x[1], x[0]))


def excel_response(applications, filename: str, title: str = 'لیست درخواست‌های پذیرش') -> HttpResponse:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    apps = list(applications)
    wb = Workbook()
    ws = wb.active
    ws.title = 'پذیرش'
    ws.append([title])
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(HEADERS))
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = Alignment(horizontal='center')
    ws.append([f'تاریخ تهیه: {format_jalali_datetime(timezone.now())} — تعداد: {len(apps)}'])
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(HEADERS))
    ws.append(HEADERS)
    for cell in ws[3]:
        cell.font = Font(bold=True)

    for row in build_rows(apps):
        ws.append(row)

    # خلاصه دسته‌بندی
    ws2 = wb.create_sheet('خلاصه')
    ws2.append(['مقطع | گروه آموزشی | رشته', 'تعداد'])
    for cell in ws2[1]:
        cell.font = Font(bold=True)
    for key, count in build_summary(apps):
        ws2.append([key, count])
    ws2.sheet_view.rightToLeft = True

    # عرض ستون‌ها بر اساس بلندترین مقدار.
    # نکته: دو ردیف نخست merge شده‌اند، پس اولین سلول هر ستون ممکن است
    # MergedCell باشد که `column_letter` ندارد و اینجا AttributeError
    # می‌داد — یعنی دکمهٔ «اکسل» همیشه ۵۰۰ می‌گرفت. حرف ستون را از
    # شمارهٔ ستون می‌گیریم و سلول‌های merge را در محاسبهٔ طول رد می‌کنیم.
    from openpyxl.utils import get_column_letter

    for index, col in enumerate(ws.columns, start=1):
        max_len = 0
        for cell in col:
            if cell.value is None:
                continue
            max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[get_column_letter(index)].width = min(max_len + 2, 36)
    ws.sheet_view.rightToLeft = True

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def word_response(applications, filename: str, title: str = 'لیست درخواست‌های پذیرش') -> HttpResponse:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    apps = list(applications)
    # جدول جمع‌وجور برای چاپ (ستون‌های کلیدی)
    print_headers = [
        'ردیف', 'کد رهگیری', 'نام و نام خانوادگی', 'کد ملی', 'موبایل',
        'مقطع', 'رشته', 'گروه', 'وضعیت',
    ]

    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        f'تاریخ تهیه: {format_jalali_datetime(timezone.now())} — تعداد: {len(apps)}'
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if meta.runs:
        meta.runs[0].font.size = Pt(10)

    summary = build_summary(apps)
    if summary:
        doc.add_heading('خلاصه بر اساس مقطع / گروه / رشته', level=2)
        st = doc.add_table(rows=1 + len(summary), cols=2)
        st.style = 'Table Grid'
        st.rows[0].cells[0].text = 'دسته'
        st.rows[0].cells[1].text = 'تعداد'
        for i, (key, count) in enumerate(summary, start=1):
            st.rows[i].cells[0].text = key
            st.rows[i].cells[1].text = str(count)

    doc.add_heading('فهرست متقاضیان', level=2)
    table = doc.add_table(rows=1 + len(apps), cols=len(print_headers))
    table.style = 'Table Grid'
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    bidi = OxmlElement('w:bidiVisual')
    bidi.set(qn('w:val'), '1')
    tblPr.append(bidi)

    for i, h in enumerate(print_headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for r_idx, app in enumerate(apps, start=1):
        major = app.desired_major
        group = ''
        if major and getattr(major, 'group_id', None) and major.group:
            group = major.group.name
        values = [
            r_idx,
            app.tracking_code or '',
            f'{app.first_name} {app.last_name}'.strip(),
            app.national_id or '',
            app.phone or '',
            app.get_degree_display() if app.degree else '',
            major.name if major else '',
            group,
            app.get_status_display() if app.status else '',
        ]
        for c_idx, value in enumerate(values):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(value)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def print_html_response(applications, title: str = 'لیست درخواست‌های پذیرش') -> HttpResponse:
    apps = list(applications)
    summary = build_summary(apps)
    rows_html = []
    for i, app in enumerate(apps, start=1):
        major = app.desired_major
        group = ''
        dept = ''
        if major:
            if getattr(major, 'group_id', None) and major.group:
                group = major.group.name
            if getattr(major, 'department_id', None) and major.department:
                dept = major.department.name
        rows_html.append(
            '<tr>'
            f'<td>{i}</td>'
            f'<td>{escape(app.tracking_code or "")}</td>'
            f'<td>{escape(f"{app.first_name} {app.last_name}".strip())}</td>'
            f'<td>{escape(app.national_id or "")}</td>'
            f'<td>{escape(app.phone or "")}</td>'
            f'<td>{escape(app.get_degree_display() if app.degree else "")}</td>'
            f'<td>{escape(major.name if major else "")}</td>'
            f'<td>{escape(group)}</td>'
            f'<td>{escape(dept)}</td>'
            f'<td>{escape(app.get_status_display() if app.status else "")}</td>'
            '</tr>'
        )

    summary_html = ''.join(
        f'<tr><td>{escape(k)}</td><td>{c}</td></tr>' for k, c in summary
    ) or '<tr><td colspan="2">موردی نیست</td></tr>'

    html = f"""<!DOCTYPE html>
<html lang="fa" dir="rtl">
<head>
<meta charset="utf-8">
<title>{escape(title)}</title>
<style>
  body {{ font-family: Tahoma, Vazirmatn, sans-serif; margin: 24px; color: #111; }}
  h1 {{ font-size: 20px; margin-bottom: 4px; }}
  .meta {{ color: #555; font-size: 13px; margin-bottom: 18px; }}
  .toolbar {{ margin-bottom: 16px; }}
  .toolbar button {{ padding: 8px 16px; font-size: 14px; cursor: pointer; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 12px; margin-bottom: 24px; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 8px; text-align: right; }}
  th {{ background: #f1f5f9; }}
  h2 {{ font-size: 15px; margin: 20px 0 8px; }}
  @media print {{
    .toolbar {{ display: none !important; }}
    body {{ margin: 8mm; }}
    a {{ text-decoration: none; color: #000; }}
  }}
</style>
</head>
<body>
  <div class="toolbar">
    <button onclick="window.print()">چاپ / ذخیره PDF</button>
  </div>
  <h1>{escape(title)}</h1>
  <div class="meta">تاریخ تهیه: {escape(format_jalali_datetime(timezone.now()))} — تعداد: {len(apps)}</div>

  <h2>خلاصه دسته‌بندی (مقطع / گروه / رشته)</h2>
  <table>
    <thead><tr><th>دسته</th><th>تعداد</th></tr></thead>
    <tbody>{summary_html}</tbody>
  </table>

  <h2>فهرست متقاضیان</h2>
  <table>
    <thead>
      <tr>
        <th>ردیف</th><th>کد رهگیری</th><th>نام</th><th>کد ملی</th><th>موبایل</th>
        <th>مقطع</th><th>رشته</th><th>گروه</th><th>دانشکده</th><th>وضعیت</th>
      </tr>
    </thead>
    <tbody>
      {''.join(rows_html) if rows_html else '<tr><td colspan="10">موردی یافت نشد</td></tr>'}
    </tbody>
  </table>
  <script>window.addEventListener('load', function() {{ /* آماده چاپ */ }});</script>
</body>
</html>"""
    return HttpResponse(html)
