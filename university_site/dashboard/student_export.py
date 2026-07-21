"""خروجی لیست دانشجویان به اکسل و ورد."""
from __future__ import annotations

from io import BytesIO

from django.http import HttpResponse
from django.utils import timezone


HEADERS = [
    'ردیف',
    'نام',
    'نام خانوادگی',
    'کد ملی',
    'شماره دانشجویی',
    'موبایل',
    'ایمیل',
    'تاریخ تولد',
    'رشته',
    'مقطع',
    'دانشکده/گروه',
    'واحد/توضیح',
    'تاریخ عضویت',
    'وضعیت حساب',
]


def student_row(index: int, profile) -> list:
    user = profile.user
    major = profile.major
    birth = ''
    if profile.birth_date:
        birth = profile.birth_date.strftime('%Y-%m-%d')
    joined = timezone.localtime(user.date_joined).strftime('%Y-%m-%d') if user.date_joined else ''
    return [
        index,
        user.first_name or '',
        user.last_name or '',
        profile.national_id or user.username or '',
        profile.student_id or '',
        profile.phone or '',
        user.email or '',
        birth,
        major.name if major else '',
        major.get_degree_display() if major else '',
        (major.group.name if major and major.group_id else '') or (major.department.name if major and major.department_id else ''),
        profile.department or '',
        joined,
        'فعال' if user.is_active else 'غیرفعال',
    ]


def build_rows(profiles) -> list[list]:
    rows = []
    for i, profile in enumerate(profiles, start=1):
        rows.append(student_row(i, profile))
    return rows


def excel_response(profiles, filename: str, title: str = 'لیست دانشجویان') -> HttpResponse:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    wb = Workbook()
    ws = wb.active
    ws.title = 'دانشجویان'
    ws.append([title])
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(HEADERS))
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = Alignment(horizontal='center')
    ws.append(HEADERS)
    for cell in ws[2]:
        cell.font = Font(bold=True)

    for row in build_rows(profiles):
        ws.append(row)

    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            val = '' if cell.value is None else str(cell.value)
            max_len = max(max_len, len(val))
        ws.column_dimensions[col_letter].width = min(max_len + 2, 40)

    ws.sheet_view.rightToLeft = True

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def word_response(profiles, filename: str, title: str = 'لیست دانشجویان') -> HttpResponse:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f'تاریخ تهیه: {timezone.localtime().strftime("%Y-%m-%d %H:%M")}')
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.runs[0].font.size = Pt(10)

    rows = build_rows(profiles)
    table = doc.add_table(rows=1 + len(rows), cols=len(HEADERS))
    table.style = 'Table Grid'

    # RTL on table
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    bidi = OxmlElement('w:bidiVisual')
    bidi.set(qn('w:val'), '1')
    tblPr.append(bidi)

    for i, h in enumerate(HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(value)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
